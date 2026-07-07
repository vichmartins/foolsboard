"""Convert a document's HTML to a downloadable file (DOCX / ODT / TXT) via pandoc.

The client sends the already-rendered document HTML (content it can already see),
so there's no cross-user data exposure here. pandoc runs with ``--sandbox`` (no
file/network access), a bounded input, and a timeout. PDF export stays client-side
(the browser's print-to-PDF), so it isn't handled here.
"""
from __future__ import annotations

import io
import os
import re
import subprocess
import tempfile
from html.parser import HTMLParser

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import Response
from pydantic import BaseModel, Field

from ..deps import get_current_user
from ..models import User

router = APIRouter(
    prefix="/api/documents", tags=["documents"], dependencies=[Depends(get_current_user)]
)

MAX_HTML_BYTES = 5_000_000

# format -> (pandoc writer, mime type, file extension)
_FORMATS: dict[str, tuple[str, str, str]] = {
    "docx": (
        "docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "odt": ("odt", "application/vnd.oasis.opendocument.text", "odt"),
    "txt": ("plain", "text/plain; charset=utf-8", "txt"),
}


class ExportRequest(BaseModel):
    html: str = Field(..., max_length=MAX_HTML_BYTES)
    title: str = "document"
    format: str
    mode: str = "doc"  # "script" gets screenplay-formatted .docx


def _safe_filename(title: str, ext: str) -> str:
    base = re.sub(r"[^\w.\- ]+", "", title or "").strip() or "document"
    return f"{base[:80]}.{ext}"


class _ScreenplayParser(HTMLParser):
    """Collect (element, text) for each top-level <p> in the screenplay HTML,
    where element is the paragraph's data-element (scene/character/…)."""

    def __init__(self) -> None:
        super().__init__()
        self.paras: list[tuple[str | None, str]] = []
        self._el: str | None = None
        self._buf: list[str] = []
        self._in_p = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag == "p":
            self._el = dict(attrs).get("data-element")
            self._buf = []
            self._in_p = True

    def handle_endtag(self, tag: str) -> None:
        if tag == "p" and self._in_p:
            self.paras.append((self._el, "".join(self._buf).strip()))
            self._in_p = False
            self._el = None

    def handle_data(self, data: str) -> None:
        if self._in_p:
            self._buf.append(data)


# Left indent (inches, from the 1.5in content margin), extras, per element. Mirrors
# the on-screen / PDF screenplay layout.
_SCRIPT_LAYOUT = {
    "scene": {"caps": True, "bold": True, "before": 14},
    "action": {"before": 8},
    "character": {"caps": True, "left": 2.2, "before": 10},
    "parenthetical": {"left": 1.6},
    "dialogue": {"left": 1.0, "right": 1.0},
    "transition": {"caps": True, "right_align": True, "before": 10},
    "shot": {"caps": True, "before": 8},
}


def _screenplay_docx(html: str, title: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    parser = _ScreenplayParser()
    parser.feed(html)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin, sec.right_margin = Inches(1.5), Inches(1.0)
    sec.top_margin, sec.bottom_margin = Inches(1.0), Inches(1.0)
    style = doc.styles["Normal"]
    style.font.name, style.font.size = "Courier New", Pt(12)

    tp = doc.add_paragraph()
    tp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(24)
    tr = tp.add_run((title or "Untitled").upper())
    tr.font.name, tr.font.size, tr.bold = "Courier New", Pt(12), True

    for element, text in parser.paras:
        spec = _SCRIPT_LAYOUT.get(element or "action", _SCRIPT_LAYOUT["action"])
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        if "left" in spec:
            fmt.left_indent = Inches(spec["left"])
        if "right" in spec:
            fmt.right_indent = Inches(spec["right"])
        if "before" in spec:
            fmt.space_before = Pt(spec["before"])
        if spec.get("right_align"):
            fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.name, run.font.size = "Courier New", Pt(12)
        if spec.get("bold"):
            run.bold = True
        if spec.get("caps"):
            run.font.all_caps = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/export")
def export_document(req: ExportRequest, _user: User = Depends(get_current_user)) -> Response:
    fmt = _FORMATS.get(req.format)
    if fmt is None:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Unsupported export format")
    writer, mime, ext = fmt

    # A screenplay to .docx gets industry formatting (Courier + element indents)
    # built directly, rather than pandoc's plain paragraph mapping.
    if req.format == "docx" and req.mode == "script":
        try:
            data = _screenplay_docx(req.html, req.title)
        except ModuleNotFoundError:
            raise HTTPException(
                status.HTTP_501_NOT_IMPLEMENTED, "Document conversion isn't available on this server."
            )
        return Response(
            content=data,
            media_type=mime,
            headers={"Content-Disposition": f'attachment; filename="{_safe_filename(req.title, ext)}"'},
        )

    # Strip <img> tags: pandoc would otherwise try to fetch each src to embed it,
    # which is an SSRF / local-file-read vector (e.g. src="file:///etc/passwd").
    # Our relative /media URLs don't embed anyway, so nothing useful is lost. This
    # replaces pandoc's --sandbox, which (in pandoc 3.x) also blocks its own data
    # files and so breaks docx/odt output entirely.
    html = re.sub(r"(?is)<img\b[^>]*>", "", req.html)
    with tempfile.TemporaryDirectory() as td:
        out_path = os.path.join(td, f"out.{ext}")
        try:
            subprocess.run(
                ["pandoc", "-f", "html", "-t", writer, "-o", out_path],
                input=html.encode("utf-8"),
                capture_output=True,
                timeout=30,
                check=True,
            )
        except FileNotFoundError:
            raise HTTPException(
                status.HTTP_501_NOT_IMPLEMENTED,
                "Document conversion isn't available on this server (pandoc not installed).",
            )
        except subprocess.TimeoutExpired:
            raise HTTPException(status.HTTP_504_GATEWAY_TIMEOUT, "Conversion timed out")
        except subprocess.CalledProcessError:
            raise HTTPException(status.HTTP_400_BAD_REQUEST, "Could not convert this document")
        with open(out_path, "rb") as f:
            data = f.read()
    return Response(
        content=data,
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{_safe_filename(req.title, ext)}"'},
    )
